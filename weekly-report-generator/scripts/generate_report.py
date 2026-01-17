#!/usr/bin/env python3
"""
周报生成器核心脚本

这个脚本的哲学：
直接生成完整的周报文档，无需外部模板。
周报的格式是固定的、可预测的，应该被编码在系统中。

设计决策：
1. 一个归类占据一行 - 建立清晰的视觉边界
2. 行内条目化展开 - 保持细节的有序性
3. 工时居中对齐 - 突出工作量分布
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys
import json


def create_weekly_report(report_data, output_path):
    """
    创建周报文档
    
    Args:
        report_data: 周报数据，格式为：
            {
                "name": "谭飞",
                "title": "后端研发",
                "date": "20251230",
                "tasks": [
                    {
                        "category": "拌合站发料管理模块",
                        "items": ["任务1", "任务2"],
                        "workload": 2.0
                    }
                ]
            }
        output_path: 输出文件路径
    """
    doc = Document()
    
    # 设置中文字体支持
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    title = doc.add_paragraph()
    title.text = f"周报_{report_data['date']}"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '姓名'
    header_cells[1].text = '职责'
    header_cells[2].text = '任务'
    header_cells[3].text = '工时(人/天)'
    
    # 设置表头格式
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    
    # 为每个任务类别添加行 - 一归类一行
    for task in report_data['tasks']:
        row_cells = table.add_row().cells
        row_cells[0].text = report_data['name']
        row_cells[1].text = report_data['title']
        
        # 构建任务文本 - 模块名 + 条目化展开
        task_text = f"{task['category']}\n"
        for i, item in enumerate(task['items'], 1):
            task_text += f"{i}. {item}\n"
        row_cells[2].text = task_text.strip()
        
        # 工时 - 智能格式化
        workload = task['workload']
        if workload == int(workload):
            row_cells[3].text = str(int(workload))
        else:
            row_cells[3].text = f"{workload:.1f}"
        
        # 工时居中对齐
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    try:
        doc.save(output_path)
        print(f"✅ 周报已成功生成：{output_path}")
    except Exception as e:
        print(f"错误：无法保存文档到 {output_path}")
        print(f"详细错误：{e}")
        sys.exit(1)


def main():
    """
    主函数 - 用于命令行调用
    
    用法：
        python generate_report.py <数据JSON文件> <输出路径>
    """
    if len(sys.argv) != 3:
        print("用法: python generate_report.py <数据JSON文件> <输出路径>")
        sys.exit(1)
    
    data_file = sys.argv[1]
    output_path = sys.argv[2]
    
    # 读取数据
    try:
        with open(data_file, 'r', encoding='utf-8') as f:
            report_data = json.load(f)
    except Exception as e:
        print(f"错误：无法读取数据文件 {data_file}")
        print(f"详细错误：{e}")
        sys.exit(1)
    
    # 生成报告
    create_weekly_report(report_data, output_path)


if __name__ == '__main__':
    main()