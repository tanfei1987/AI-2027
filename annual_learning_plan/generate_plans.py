from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_weekly_plan(doc_path, title, weeks_data):
    doc = Document()
    
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('学习周期：按周拆分计划')
    doc.add_paragraph()
    
    table = doc.add_table(rows=len(weeks_data)+1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['周次', '学习主题', '具体内容', '学习资源', '主要概念']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    for row_idx, week_data in enumerate(weeks_data):
        for col_idx, value in enumerate(week_data):
            table.rows[row_idx+1].cells[col_idx].text = str(value)
    
    doc.save(doc_path)
    print(f'已生成：{doc_path}')

# 1. 工程项目整体管理
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/01_工程项目整体管理_周计划.docx', '工程项目整体管理 - 周学习计划', [
    ['第1周', '项目全生命周期管理', '学习项目启动、规划、执行、监控、收尾5个阶段；绘制项目生命周期流程图；阅读PMBOK指南第1-2章', '《PMBOK指南》第7版（PMI官网下载）https://www.pmi.org/pmbok-guide-standards\n《项目管理知识体系指南》', '项目生命周期、项目阶段、项目管理过程组'],
    ['第2周', '工程项目 vs 工业项目', '对比工程项目与工业项目的特点差异；分析工程项目的独特性；收集3个实际工程项目案例', '《工程项目管理》同济大学出版社\n《工程项目管理》高校教材（第3版）\n中国工程项目管理网 http://www.cpmchina.com', '工程项目特点、工业项目差异、独特性、临时性'],
    ['第3周', '项目管理组织架构', '学习项目管理组织的三种模式；分析各模式的优缺点；绘制组织架构图示例', '《建设工程项目管理规范》GB/T 50326-2017\n《项目管理组织设计》机械工业出版社\n住建部官网 http://www.mohurd.gov.cn', '职能型组织、矩阵型组织、项目型组织、组织架构'],
    ['第4周', '综合实践', '完成项目管理模拟案例；总结项目整体管理要点；制作思维导图', '《PMBOK指南》附录A（术语表）\n《项目管理案例集》中国电力出版社\nXMind思维导图 https://xmind.cn', '项目整体管理要点、案例分析、经验总结']
])

# 2. 项目物资采购管理
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/02_项目物资采购管理_周计划.docx', '项目物资采购管理 - 周学习计划', [
    ['第1周', '采购与招标基础', '学习采购的基本流程和分类；掌握招标方式；阅读招标投标法要点', '《招标投标法》 http://www.ndrc.gov.cn\n《政府采购法》财政部官网\n中国招标公共服务平台 http://www.cebpubservice.com', '招标、投标、公开招标、邀请招标、采购流程'],
    ['第2周', '甲供/乙供/材料分类', '区分甲供材、乙供材范围；制作材料分类清单模板；分析材料供应模式案例', '《建设工程施工合同（示范文本）》GF-2017-0201\n《建筑工程物资管理》化学工业出版社\n筑龙网 https://www.zhulong.com', '甲供材、乙供材、材料分类、供应模式'],
    ['第3周', '采购流程', '梳理采购6个阶段；绘制采购流程图；学习供应商管理方法', '《采购管理实务》中国人民大学出版社\n《供应商管理》机械工业出版社\n采购管理师教材 http://www.cpsm.org.cn', '采购需求、供应商选择、合同签订、供应商管理'],
    ['第4周', '设备管理', '学习设备选型和采购要点；掌握设备验收流程；制作设备采购清单', '《机电工程管理与实务》中国建筑工业出版社\n《设备采购与安装工程》化学工业出版社\n设备验收规范 http://www.cseps.com', '设备选型、设备采购、设备验收、设备清单']
])

# 3. 项目进度计划编制
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/03_项目进度计划编制_周计划.docx', '项目进度计划编制 - 周学习计划', [
    ['第1周', 'WBS工作分解结构', '学习WBS分解原则；掌握WBS分解方法；练习绘制WBS分解图', '《PMBOK指南》第6版第5章 https://www.pmi.org\n《项目进度管理》中国电力出版社\nWBS教程 https://www.projectmanager.com/wbs', 'WBS、工作分解结构、100%原则、层级原则'],
    ['第2周', '里程碑与关键路径', '学习里程碑设置方法；掌握关键路径法原理；计算关键路径和浮动时间', '《网络计划技术》冶金工业出版社\n《关键路径法CPM》教程\n进度管理师教材 http://www.miit.gov.cn', '里程碑、关键路径、关键活动、浮动时间'],
    ['第3周', 'Project软件实操', '安装和熟悉Microsoft Project；编制项目进度计划；设置里程碑和关键路径', 'Microsoft Project https://www.microsoft.com/zh-cn/microsoft-365/project\n《Microsoft Project从入门到精通》人民邮电出版社\nB站教程 https://www.bilibili.com', 'Project操作、甘特图、里程碑设置、资源分配'],
    ['第4周', '进度偏差分析', '学习进度偏差分析方法；掌握进度调整措施；模拟进度延误情景并调整', '《项目进度控制》中国建筑工业出版社\n《挣值管理》PMI官方指南\nS曲线分析 https://www.pmi.org/learning/earned-value-management', '前锋线、S曲线、进度偏差、进度调整']
])

# 4. 合同/成本/进度控制
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/04_合同成本进度控制_周计划.docx', '合同/成本/进度控制 - 周学习计划', [
    ['第1周', '进度控制方法', '学习进度控制的基本原理；掌握进度监控的方法和工具；学习进度预警机制', '《项目进度控制》中国建筑工业出版社\n《挣值管理》中国电力出版社\nPMI进度管理 https://www.pmi.org/learning', '进度监控、挣值管理、进度绩效指数、预警机制'],
    ['第2周', '合同管理', '学习合同条款解读；掌握合同变更管理流程；分析合同风险点', '《合同法》 http://www.npc.gov.cn\n《建设工程施工合同（示范文本）》\n《FIDIC合同条件》 http://www.fidic.org', '合同条款、合同变更、合同风险、合同管理'],
    ['第3周', '变更与索赔', '学习变更管理的程序和要求；掌握索赔依据和证据收集；模拟索赔案例分析', '《工程索赔》中国建筑工业出版社\n《索赔与反索赔》机械工业出版社\n《建设工程施工合同司法解释》', '工程变更、索赔程序、索赔依据、证据收集'],
    ['第4周', '纠纷处理', '学习工程纠纷处理程序；掌握调解、仲裁、诉讼流程；分析典型纠纷案例', '《工程纠纷案例精选》法律出版社\n《仲裁法》中国法律网 http://www.court.gov.cn\n《中国建设工程法律评论》', '工程纠纷、调解、仲裁、诉讼、争议解决']
])

# 5. 施工管理
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/05_施工管理_周计划.docx', '施工管理 - 周学习计划', [
    ['第1周', '施组设计', '学习施工组织设计的内容和编制要求；掌握施工部署和施工方案制定；编制简单工程的施组设计', '《建筑施工组织设计规范》GB/T 50502-2009\n《建筑施工手册》第5版 中国建筑工业出版社\n筑龙网 https://www.zhulong.com/dict', '施工部署、施工方案、施组设计、工艺流程'],
    ['第2周', '安装调试', '学习设备安装工艺流程；掌握调试方法和验收标准；了解安装过程的质量控制', '《机电工程管理与实务》中国建筑工业出版社\n《工业设备安装工程施工规范》GB 50231\n设备安装手册 http://www.cbichina.org.cn', '设备安装、工艺流程、调试方法、验收标准'],
    ['第3周', '现场管理', '学习现场管理的要素；掌握现场布置原则；制定现场管理检查表', '《施工现场管理》中国建筑工业出版社\n《安全文明施工标准化》中国电力出版社\n5S管理协会 http://www.sechina.org', '人机料法环、现场布置、安全文明、5S管理'],
    ['第4周', '综合实践', '参与现场巡查实践；总结施工管理要点；制作施工管理手册', '《建筑施工安全检查标准》JGJ 59-2011\n《工程项目管理》案例教程\nBIM技术 http://www.cbim.org.cn', '现场巡查、质量控制、安全管理、进度协调']
])

# 6. 项目现场管理
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/06_项目现场管理_周计划.docx', '项目现场管理 - 周学习计划', [
    ['第1周', '项目部管理', '学习项目部组织架构设置；掌握项目部管理制度制定；明确岗位职责和权限', '《建设工程项目管理规范》GB/T 50326\n《项目部设置与管理》中国建筑工业出版社\n建造师教材 http://www.coc.gov.cn', '项目部架构、管理制度、岗位职责、权限分配'],
    ['第2周', '施工现场管理', '学习施工现场布置原则；掌握施工平面图绘制；制定现场安全文明施工措施', '《建筑施工现场管理》化学工业出版社\n《施工平面图绘制》教程\n安全文明施工标准 http://www.mohurd.gov.cn', '施工平面图、现场布置、安全文明、施工道路'],
    ['第3周', '仓储管理', '学习仓储管理流程；掌握物资入库、出库、盘点；优化仓储布局提高效率', '《仓储管理实务》中国人民大学出版社\n《物资管理手册》机械工业出版社\nERP系统学习 https://www.netsuite.com', '仓储流程、入库出库、盘点、物资管理'],
    ['第4周', '预安装管理', '学习预安装范围确定；掌握预安装计划编制；制定预安装质量控制措施', '《装配式混凝土建筑技术标准》GB/T 51231\n《预制构件安装技术规程》\n装配式建筑学习网 http://www.cnieia.org', '预安装范围、预安装计划、质量控制、成品保护']
])

# 7. 施工技术与发展
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/07_施工技术与发展_周计划.docx', '施工技术与发展 - 周学习计划', [
    ['第1周', '行业技术标准', '收集常用施工技术标准；学习标准规范查询方法；整理关键技术标准要点', '国家标准全文公开系统 http://www.gb688.cn\n《工程建设标准化》协会官网 http://www.ciecc.org.cn\n住建部标准定额司 http://www.mohurd.gov.cn', '国家标准、行业标准、企业标准、规范体系'],
    ['第2周', '规范与图集', '学习施工规范体系；掌握常用图集查阅；规范与图集的实际应用', '国标图集 中国建筑标准设计研究院 https://www.china-building.com.cn\n《混凝土结构施工图集》16G101\n地方标准图集 http://www.cabplink.com', '施工规范、图集查阅、构造做法、标准图集'],
    ['第3周', '施工技术要点', '学习各专业施工技术要点；掌握关键技术工艺流程；总结常见技术问题解决方法', '《建筑施工技术》同济大学出版社\n《施工工艺手册》化学工业出版社\n工程技术交底模板 http://www.cbi360.net', '施工工艺、技术要点、质量标准、工艺流程'],
    ['第4周', '技术发展趋势', '了解装配式建筑技术；学习BIM技术在施工中的应用；关注绿色施工技术发展', '《装配式建筑概论》中国建筑工业出版社\n《BIM技术应用》中国建筑工业出版社\n《绿色建筑评价标准》GB/T 50378', '装配式建筑、BIM技术、绿色施工、智能建造']
])

# 8. 成本与造价管理
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/08_成本与造价管理_周计划.docx', '成本与造价管理 - 周学习计划', [
    ['第1周', '预算编制', '学习施工图预算编制方法；掌握工程量清单计价；练习使用造价软件', '《建设工程工程量清单计价规范》GB 50500-2013\n《施工图预算》中国建筑工业出版社\n广联达造价 https://www.glodon.com\n品茗造价 https://www.pmddjj.com', '施工图预算、工程量清单、清单计价、定额'],
    ['第2周', '成本核算', '学习成本核算的对象和方法；掌握成本归集和分配；编制成本报表', '《施工企业会计》中国建筑工业出版社\n《成本管理》中国人民大学出版社\n《会计准则第15号》财政部 http://kjs.mof.gov.cn', '成本核算对象、成本归集、成本分配、成本报表'],
    ['第3周', '成本控制体系', '建立成本控制体系框架；制定成本控制措施；学习目标成本管理', '《工程项目成本控制》机械工业出版社\n《目标成本管理》企业管理出版社\n《精益建造》中国建筑工业出版社', '目标成本、成本控制体系、责任成本、成本考核'],
    ['第4周', '成本分析', '学习成本偏差分析方法；掌握成本纠偏措施；编制成本分析报告', '《成本分析》中国电力出版社\n《工程成本管理》化学工业出版社\n造价工程师教材 http://www.ceca.org.cn', '成本偏差、纠偏措施、成本分析、成本报告']
])

# 9. EPC总承包项目管理
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/09_EPC总承包项目管理_周计划.docx', 'EPC总承包项目管理 - 周学习计划', [
    ['第1周', 'EPC模式特点', '学习EPC总承包模式概念；对比传统承包模式差异；分析EPC模式优劣势', '《房屋建筑和市政基础设施项目工程总承包管理办法》\n《EPC工程总承包项目管理》中国建筑工业出版社\n总承包管理网 http://www.epc365.com', 'EPC总承包、设计采购施工、施工总承包、DBB模式'],
    ['第2周', '设计采购施工协调', '学习设计与采购接口管理；掌握设计与施工协调要点；优化设计-采购-施工流程', '《工程总承包项目管理》机械工业出版社\n《接口管理》中国电力出版社\n设计管理手册 http://www.chinaeda.org', '设计-采购接口、设计-施工协调、接口管理、协同流程'],
    ['第3周', '项目总控', '学习项目总控理念和方法；掌握项目总控组织架构；建立项目总控信息系统', '《项目总控管理》中国建筑工业出版社\n《PMC项目管理》化学工业出版社\n项目控制系统 http://www.pmshop.cn', '项目总控、PMC、项目控制、组织架构'],
    ['第4周', '案例分析', '分析EPC项目成功案例；总结EPC项目风险点；制定EPC项目管理要点', '《EPC工程总承包案例分析》中国建筑工业出版社\n《国际工程承包》案例集\nFIDIC合同应用 http://www.fidic.org.cn', 'EPC风险、案例分析、最佳实践、管理要点']
])

# 10. 多Agent系统
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/10_多Agent系统_周计划.docx', '多Agent系统 - 周学习计划', [
    ['第1周', 'Agent架构基础', '学习Agent的基本概念和架构；了解Planner/Executor/Critic模式；分析Agent应用场景', 'LangChain官方文档 https://python.langchain.com\n吴恩达《AI Agent》课程 https://www.deeplearning.ai\nAutoGen官方文档 https://microsoft.github.io/autogen', 'Agent、Planner、Executor、Critic、协作机制'],
    ['第2周', '多Agent协作机制', '学习多Agent通信协议；掌握Agent协作方式；设计简单的Agent交互流程', '《Multi-Agent Systems》论文集\nCrewAI官方文档 https://crewai.com\nMetaGPT GitHub https://github.com/geekan/MetaGPT', 'Agent通信、协作协议、任务分配、交互流程'],
    ['第3周', 'Agent框架实践', '了解主流Agent框架；搭建简单的多Agent系统；实现Agent间的任务分配', 'LangChain Agents https://python.langchain.com/docs/modules/agents\nAutoGen示例库 https://github.com/microsoft/autogen\nLlamaIndex Agents https://docs.llamaindex.ai', 'LangChain、AutoGen、CrewAI、任务分配'],
    ['第4周', '综合应用', '完成多Agent系统小项目；优化Agent协作效率；总结多Agent系统设计要点', 'GitHub Agent项目 https://github.com/topics/agent\nHugging Face Agents https://huggingface.co/agents\nAI Agent实战案例 https://www.langchain.com/agents', '系统设计、性能优化、经验总结、架构设计']
])

# 11. 编程辅助工具
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/11_编程辅助工具_周计划.docx', '编程辅助工具 - 周学习计划', [
    ['第1周', 'OpenCode基础', '安装和配置OpenCode；学习基本命令和操作；完成第一个OpenCode项目', 'OpenCode官网 https://opencode.ai\nOpenCode GitHub https://github.com/anomalyco/opencode\nOpenCode文档 https://docs.opencode.ai', 'OpenCode安装、项目创建、命令操作、代码生成'],
    ['第2周', 'OpenCode进阶', '掌握代码生成功能；学习项目管理和协作；使用OpenCode重构代码', 'OpenCode最佳实践 https://docs.opencode.ai/guides\n《代码重构》Martin Fowler 机械工业出版社\n代码质量指南 https://google.github.io/styleguide', '代码生成、重构、项目管理、多文件协作'],
    ['第3周', 'Claude Code应用', '学习Claude Code安装配置；掌握代码分析和优化；使用Claude Code辅助调试', 'Claude Code https://claude.com/claude-code\nAnthropic开发者文档 https://docs.anthropic.com\nClaude使用指南 https://www.anthropic.com/claude', 'Claude Code安装、代码分析、调试辅助、代码审查'],
    ['第4周', '工具整合', '整合OpenCode和Claude Code；建立高效编程工作流；总结工具使用技巧', 'AI编程工具评测 https://www.toolify.ai\n《高效程序员》 O Reilly出版社\n工作流自动化 https://n8n.io', '工具整合、工作流优化、效率提升、最佳实践']
])

# 12. 大模型原理
create_weekly_plan('C:/Users/admin/.config/opencode/skills/annual_learning_plan/detail/12_大模型原理_周计划.docx', '大模型原理 - 周学习计划', [
    ['第1周', 'Transformer架构', '学习Attention注意力机制；掌握Transformer模型结构；解读Transformer原论文', '《Attention Is All You Need》论文 https://arxiv.org/abs/1706.03762\n李宏毅Transformer课程 https://www.youtube.com/watch?v=nFBqBlM5mvc\n《深度学习》花书 第11章', 'Self-Attention、Multi-Head Attention、位置编码、Encoder-Decoder'],
    ['第2周', 'MoE混合专家模型', '学习MoE架构原理；掌握稀疏激活机制；分析MoE模型优势', 'Switch Transformer论文 https://arxiv.org/abs/2101.03961\nGShard论文 https://arxiv.org/abs/2006.16668\nMoE教程 https://www.vwalsh.com/posts/mixture-of-experts', '混合专家模型、稀疏激活、门控机制、负载均衡'],
    ['第3周', 'SFT监督微调', '学习SFT训练方法；掌握数据准备和标注；实践SFT微调流程', 'LLaMA-Factory GitHub https://github.com/hiyouga/LLaMA-Factory\nAlpaca数据集 https://github.com/tatsu-lab/stanford_alpaca\nSFT教程 https://www.philschmid.de/fine-tuning-llamas', '监督微调、数据标注、指令微调、损失函数'],
    ['第4周', 'RLHF强化学习', '学习RLHF训练框架；掌握PPO算法原理；了解Reward Model训练', 'InstructGPT论文 https://arxiv.org/abs/2203.02155\nPPO算法详解 https://spinningup.openai.com/en/algorithms/ppo.html\nRLHF教程 https://huggingface.co/blog/rlhf', 'RLHF、PPO算法、Reward Model、策略梯度'],
    ['第5周', '开源模型实践', '学习使用LLAMA等开源模型；掌握模型部署方法；实践模型微调和推理', 'Hugging Face https://huggingface.co\nLLAMA模型 https://llama.meta.com\nOllama本地部署 https://ollama.com\nGGUF模型 https://github.com/ggerganov/llama.cpp', '开源模型、模型部署、模型微调、推理优化']
])

print('已生成12个周计划文档（含书籍信息和下载链接）！')
